import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime

st.set_page_config(page_title="Certificados UCJC", layout="centered")
st.title("🎓 Generador de Certificados Docentes UCJC")

# Subir archivos
docencia_file = st.file_uploader("📚 Sube el Excel de Docencia", type=["xlsx"])
tfm_file = st.file_uploader("📄 Sube el Excel de TFMs", type=["xlsx"])
plantilla_file = st.file_uploader("📁 Sube la plantilla Word", type=["docx"])

nombre_profesor = st.text_input("👨‍🏫 Nombre completo del profesor (como en el Excel)")

if st.button("✅ Generar Certificado") and all([docencia_file, tfm_file, plantilla_file, nombre_profesor]):
    df_docencia = pd.read_excel(docencia_file)
    df_tfm = pd.read_excel(tfm_file)
    doc = Document(plantilla_file)

    asignaturas = df_docencia[df_docencia["Firmante"] == nombre_profesor][
        ["Estudios", "C.Asg", "Asignatura", "Crd"]
    ].drop_duplicates()

    tfms = df_tfm[df_tfm["Tutor asignado"] == nombre_profesor][
        ["Estudios"]
    ].value_counts().reset_index(name='N_TFM')
    tfms["Curso"] = "2024-2025"
    tfms["ECTS"] = 12
    tfms["Obs"] = "En defensa"

    # Insertar asignaturas
    tabla_docencia = doc.tables[0]
    for _, row in asignaturas.iterrows():
        c = tabla_docencia.add_row().cells
        c[0].text = row["Estudios"]
        c[1].text = str(row["C.Asg"])
        c[2].text = row["Asignatura"]
        c[3].text = str(row["Crd"])
        c[4].text = "15"
        c[5].text = "Teórica"

    # Insertar TFMs
    tabla_tfm = doc.tables[1]
    for _, row in tfms.iterrows():
        c = tabla_tfm.add_row().cells
        c[0].text = row["Estudios"]
        c[1].text = str(row["N_TFM"])
        c[2].text = row["Curso"]
        c[3].text = str(row["ECTS"])
        c[4].text = row["Obs"]

    # Sustituir textos
    for p in doc.paragraphs:
        if "[Nombre del Profesor/a]" in p.text:
            p.text = p.text.replace("[Nombre del Profesor/a]", nombre_profesor)
        if "[XXXXXXXXX]" in p.text:
            p.text = p.text.replace("[XXXXXXXXX]", "[DNI]")
        if "__de ___ de 2025" in p.text:
            fecha = datetime.now().strftime("%d de %B de %Y")
            p.text = p.text.replace("__de ___ de 2025", fecha)

    nombre_salida = f"certificado_{nombre_profesor.replace(' ', '_').replace(',', '')}.docx"
    doc.save(nombre_salida)

    with open(nombre_salida, "rb") as file:
        st.download_button("📥 Descargar certificado", file, file_name=nombre_salida)
