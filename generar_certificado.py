import pandas as pd
from docx import Document
from datetime import datetime

# === CONFIGURACIÓN ===
archivo_docencia = "docencia.xlsx"            # Tu Excel con asignaturas
archivo_tfms = "tfms.xlsx"                    # Tu Excel con TFMs
plantilla_word = "certificado_tutorizacion.docx"
nombre_profesor = input("Introduce el nombre completo del profesor: ")

# === CARGA DE DATOS ===
df_docencia = pd.read_excel(archivo_docencia)
df_tfms = pd.read_excel(archivo_tfms)

# === FILTROS ===
asignaturas = df_docencia[df_docencia["Firmante"] == nombre_profesor][
    ["Estudios", "C.Asg", "Asignatura", "Crd"]
].drop_duplicates()

tfms = df_tfms[df_tfms["Tutor asignado"] == nombre_profesor][
    ["Estudios"]
].value_counts().reset_index(name='N_TFM')
tfms["Curso"] = "2024-2025"
tfms["ECTS"] = 12
tfms["Obs"] = "En defensa"

# === CARGA PLANTILLA ===
doc = Document(plantilla_word)

# === RELLENO DE TABLA DE DOCENCIA ===
tabla_docencia = doc.tables[0]

for _, row in asignaturas.iterrows():
    nueva_fila = tabla_docencia.add_row().cells
    nueva_fila[0].text = row["Estudios"]
    nueva_fila[1].text = str(row["C.Asg"])
    nueva_fila[2].text = row["Asignatura"]
    nueva_fila[3].text = str(row["Crd"])
    nueva_fila[4].text = "15"          # O personaliza según Excel
    nueva_fila[5].text = "Teórica"     # O calcula según naturaleza

# === RELLENO DE TABLA DE TFM ===
tabla_tfm = doc.tables[1]

for _, row in tfms.iterrows():
    nueva_fila = tabla_tfm.add_row().cells
    nueva_fila[0].text = row["Estudios"]
    nueva_fila[1].text = str(row["N_TFM"])
    nueva_fila[2].text = row["Curso"]
    nueva_fila[3].text = str(row["ECTS"])
    nueva_fila[4].text = row["Obs"]

# === RELLENO DE NOMBRE Y FECHA EN TEXTO ===
for p in doc.paragraphs:
    if "[Nombre del Profesor/a]" in p.text:
        p.text = p.text.replace("[Nombre del Profesor/a]", nombre_profesor)
    if "[XXXXXXXXX]" in p.text:
        p.text = p.text.replace("[XXXXXXXXX]", "[DNI POR COMPLETAR]")  # Puedes automatizar si lo tienes
    if "__de ___ de 2025" in p.text:
        fecha_hoy = datetime.now().strftime("%d de %B de %Y")
        p.text = p.text.replace("__de ___ de 2025", fecha_hoy)

# === GUARDAR ===
nombre_archivo = f"certificado_{nombre_profesor.replace(' ', '_').replace(',', '')}.docx"
doc.save(nombre_archivo)
print(f"✅ Certificado generado: {nombre_archivo}")
